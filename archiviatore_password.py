#Programma che preleva una password casuale e la inserisce in una tabella su word con due colonne una per la pswd e una per il sito 
from docx import Document
from docx.shared import Cm
import random

document = Document()

# Programma genera e salva password
possibili_caratteri = ''
numero_caratteri = ''
password = ''

# Richiedi se vuole lettere minuscole
def richiesta_lettere_minuscole() :

	scelta_password = ''
	print('''Vuoi le lettere minuscole? [si/no]''')
	scelta_password = str(input())
	
	while((scelta_password != 'si') and (scelta_password != 'no') and (scelta_password != 'sì')):
		
		print('Errore')
		print('''Vuoi le lettere minuscole? [si/no]''')
		scelta_password = str(input())
	return scelta_password

# Richiedi se vuole lettere maiuscole
def richiesta_lettere_maiuscole() :

	scelta_password = ''
	print('''Vuoi le lettere maiuscole? [si/no]''')
	scelta_password = str(input())
	
	while((scelta_password != 'si') and (scelta_password != 'no') and (scelta_password != 'sì')):
		
		print('Errore')
		print('''Vuoi le lettere maiuscole? [si/no]''')
		scelta_password = str(input())
	return scelta_password

# Richiedi se vuole numeri
def richiesta_numeri() :

	scelta_password = ''
	print('''Vuoi i numeri? [si/no]''')
	scelta_password = str(input())
	
	while((scelta_password != 'si') and (scelta_password != 'no') and (scelta_password != 'sì')):
		
		print('Errore')
		print('''Vuoi i numeri? [si/no]''')
		scelta_password = str(input())
	return scelta_password

# Richiedi se vuole i caratteri ascii
def richiesta_ascii() :

	scelta_password = ''
	print('''Vuoi i caratteri ascii? [si/no]''')
	scelta_password = str(input())
	
	while((scelta_password != 'si') and (scelta_password != 'no') and (scelta_password != 'sì')):
		
		print('Errore')
		print('''Vuoi i caratteri ascii? [si/no]''')
		scelta_password = str(input())
	return scelta_password

stringa_lettere_maiuscole = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
stringa_lettere_minuscole = 'abcdefghijklmnopqrstuvwxyz'
stringa_numeri = '0987654321'
stringa_ascii = " !#£$%&()*+-/:;=<>?@"

# Richiesta e controllo lunghezza password
while(numero_caratteri == ''):
	numero_caratteri = int(input('Inserisci la lunghezza della password (devono essere almeno 5 caratteri): '))
	if(numero_caratteri < 4):
		print('Password troppo corta:')
		numero_caratteri = ''

	if(numero_caratteri > 10000):
		print('Password troppo lunga:')
		numero_caratteri = ''

# Se selezionato aggiunta lettere minuscole ai caratteri da usare nella password
if((richiesta_lettere_minuscole()) != 'no'):
	possibili_caratteri = possibili_caratteri + stringa_lettere_minuscole
	# Aggiunge a password almeno una lettera minuscola se sono volute nella password
	password = password + (random.choice(stringa_lettere_minuscole))
	numero_caratteri = numero_caratteri - 1

# Se selezionato aggiunta lettere maiuscole ai caratteri da usare nella password
if((richiesta_lettere_maiuscole()) != 'no'):
	possibili_caratteri = possibili_caratteri + stringa_lettere_maiuscole
	# Aggiunge a password almeno una lettera maiuscola se sono volute nella password
	password = password + (random.choice(stringa_lettere_maiuscole))
	numero_caratteri = numero_caratteri - 1

# Se selezionato aggiunta numeri ai caratteri da usare nella password
if((richiesta_numeri()) != 'no'):
	possibili_caratteri = possibili_caratteri + stringa_numeri
	# Aggiunge a password almeno un numero se sono voluti nella password
	password = password + (random.choice(stringa_numeri))
	numero_caratteri = numero_caratteri - 1

# Se selezionato aggiunta caratteri ascii ai caratteri da usare nella password
if((richiesta_ascii()) != 'no'):
	possibili_caratteri = possibili_caratteri + stringa_ascii
	# Aggiunge a password almeno un carattere ascii se sono voluti nella password
	password = password + (random.choice(stringa_ascii))
	numero_caratteri = numero_caratteri - 1

# Generazione password
for x in range(numero_caratteri):
	password = password + (random.choice(possibili_caratteri))

# Rimescolamento password
password = ''.join(random.sample(password,len(password)))
print('La tua password è: ')
print(password)

nuovo_sito = ''
nuovo_sito = str(input('Inserisci il nome del nuovo sito a cui vuoi collegare la Password: '))
		
#Titolo: Archiviatore Password
document.add_heading('Archiviatore password',0)

item_archiviatore = [
				('Gmail', 'OsWaLdO'),
				('Instagram', '1q2w3e4r'),
				('Facebook', 'LaPizZaChePiace96'),
				('Twitter', 'LaQuiLa')
					]

nuova_password = password 
nuovi_elementi = (nuovo_sito, nuova_password)
item_archiviatore.insert(0, nuovi_elementi)

#Creiamo la tabella 'tabella_archiviatore'
tabella_archiviatore = document.add_table(rows=1, cols=2)
#Stile a griglia per la tabella creata
tabella_archiviatore.style = 'TableGrid'
#Creiamo la Riga Superiore
heading_cells = tabella_archiviatore.rows[0].cells
heading_cells[0].text = 'Sito:'
heading_cells[1].text = 'Password:'
#Creiamo righe e celle per ciascun elemento nella nostra lista 'item_archiviatore'
for item in item_archiviatore:
	celle_archiviatore = tabella_archiviatore.add_row().cells
	celle_archiviatore[0].text = str(item[0])
	celle_archiviatore[1].text = str(item[1])

document.save('archiviatore_password.docx')